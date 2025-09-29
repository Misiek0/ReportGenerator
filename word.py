import os
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT


def open_docx(filename):
    curr_dir = os.getcwd()
    filepath = os.path.join(curr_dir,"templates",filename)
    opened_doc = Document(filepath)
    return opened_doc

def save_docx(save_name,file):
    curr_dir = os.getcwd()
    filepath = os.path.join(curr_dir,"reports",save_name)
    os.makedirs(os.path.dirname(filepath), exist_ok=True)
    file.save(filepath)

def find_col_index(target_colname, table):
    header_row = table.rows[0]
    header_cells = [cell.text.strip() for cell in header_row.cells]
    if target_colname not in header_cells:
        raise ValueError(f"Nie znaleziono kolumny o nazwie {target_colname}")
    col_index = header_cells.index(target_colname)
    return col_index

def find_row_index(target_automat_id, table):
    col_index = find_col_index("Nr seryjny\nautomatu biletowego", table)
    for index, row in enumerate(table.rows):
        cell_value = row.cells[col_index].text.strip()
        if target_automat_id == cell_value:
            return index
    return None

def format_paragraph(paragraph,isBold):
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    for run in paragraph.runs:
        run.font.name = 'Arial'
        run.font.size = Pt(8)
        run.bold = isBold

def insert_failures(automat_id, col_index, failures_dict, table, solution_dict):
    if 'EN' in automat_id:
        word_automat_id = automat_id[0:4]
    else:
        id_localization_list = automat_id.split()
        word_automat_id = str(id_localization_list[0].strip())
    row_index = find_row_index(word_automat_id, table)

    if row_index is None:
        raise ValueError(f"Zwrócono pusty indeks {row_index}")
    cell = table.cell(row_index, col_index)
    values = [f"{count}x {solution_dict[failure]}" for failure,count in failures_dict[automat_id].items()]

    if cell.text:  # jeżeli w komórce już coś jest
        cell.text += ", " + ", ".join(values)
    else:  # jeżeli komórka pusta, to nie dodawaj przecinka na początku
        cell.text = ", ".join(values)

    for paragraph in cell.paragraphs:
        format_paragraph(paragraph,False)

def replace_text(doc, replacement_dict):
    for paragraph in doc.paragraphs:
        text = paragraph.text
        if '{month}' in text or '{year}' in text:
            paragraph.text = text.replace('{month}', replacement_dict['{month}']).replace('{year}', replacement_dict['{year}'])
        format_paragraph(paragraph, True)
